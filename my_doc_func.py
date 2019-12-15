#Гребанные сонники
from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
import pymorphy2
from tokenize_functions import *


decision_value=0.75

def read_file_to_doc_sonic(filename):
    count_change=0
    with open('Lyrics\\'+filename+'.txt',encoding='utf-8') as f:
        document=Document()
        for line in f:
            was_symbol=0
            if len(line)>1:
                paragraph=document.add_paragraph('')
                prev_word=['',False]
                for word in line.split():
                    if was_symbol>0:
                        prev_word[1]=True
                    else:
                        prev_word[1]=False
                    was_symbol=0
                    run = paragraph.add_run(word+' ')
                    tokens=prepare_text_for_lda(word)
                    p = morph.parse(word)[0]
                    if len(tokens)>0:
                        if (tokens[0] not in useless_set):
                            for sonic_symbol in sonic_symbols:
                                if tokens[0] in sonic_symbol:
                                    if prev_word[1]:
                                        run.font.highlight_color=4
                                    count_change+=1
                                    was_symbol+=1
                                    break
                                    
                            prev_word[0]=tokens[0]
                                
        document.save('Outputs//test_Sonic.docx')

def check_ngramm(ngramm,decision_value):
    count=0
    res=decision_value*len(ngramm)
    for gramm in ngramm:
        tokens=prepare_text_for_lda(gramm)
        if not(len(tokens)>0 and (tokens[0] not in useless_set)):
            continue
        for symbol in symbols:
            if tokens[0] in symbol:
                count+=1
                break
    if (count>=res)and len(ngramm)>1:
        return True
    else:
        return False
    #print('{} = {}'.format(str(count),str(res)))
    

def read_file_to_doc_ngramm(filename,len_ngramm,decision_value):
    #ngramm=list(n)
    count_change=0
    count_symbols_var=0
    prev_value=False
    ngramm=[]
    with open('Lyrics\\'+filename+'.txt',encoding='utf-8') as f:
        document=Document()
        for line in f:
            prev_value=False
            if not len(line)>1:
                continue
            temp_bool=check_ngramm(ngramm,decision_value)
            for gramm in ngramm:
                run = paragraph.add_run(gramm+' ')
                if temp_bool:
                    run.font.highlight_color=4
                else:
                    run.font.highlight_color=None
            ngramm=[]
            paragraph=document.add_paragraph('')
            for word in line.split():
                if len(ngramm)==len_ngramm:
                    temp_bool=check_ngramm(ngramm,decision_value)
                    run = paragraph.add_run(ngramm.pop(0)+' ')
                    if temp_bool or prev_value:
                        run.font.highlight_color=4
                    else:
                        run.font.highlight_color=None
                    prev_value=temp_bool
                ngramm.append(word)
                
        temp_bool=check_ngramm(ngramm,decision_value)
        for gramm in ngramm:
            run = paragraph.add_run(gramm+' ')
            if temp_bool:
                run.font.highlight_color=4
            else:
                run.font.highlight_color=None
        document.save('Outputs//test_ngramm'+str(len_ngramm)+'.docx')